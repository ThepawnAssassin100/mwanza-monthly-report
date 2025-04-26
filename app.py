import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode

def generate_word_report(dataframes, officer_name, report_month):
    doc = Document()
    doc.add_heading('MINISTRY OF AGRICULTURE', 0)
    doc.add_heading('DEPARTMENT OF IRRIGATION', level=1)
    doc.add_heading('MWANZA DISTRICT IRRIGATION OFFICE', level=1)
    doc.add_heading('MONTHLY PROGRESS REVIEW REPORT', level=2)
    doc.add_paragraph(f'({report_month.strftime("%B %Y")})')

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("Summary of the activities and progress for the reporting month.")

    doc.add_heading('Contents', level=1)
    doc.add_paragraph("Auto-generated during final editing.")

    sections = [
        ("Introduction", "This report provides an overview of irrigation activities..."),
        ("Human, Financial and Motor Vehicle Resources", dataframes['staffing']),
        ("Vehicles Status", dataframes['vehicles']),
        ("ORT Budget Performance", dataframes['budget']),
        ("Physical Implementation Progress", dataframes['development']),
        ("Scheme Utilisation", dataframes['utilisation']),
        ("Collaboration with Other Stakeholders", dataframes['stakeholders']),
        ("Challenges", dataframes['challenges']),
        ("Recommendations", dataframes['recommendations'])
    ]

    for title, content in sections:
        doc.add_heading(title, level=2)
        if isinstance(content, pd.DataFrame):
            table = doc.add_table(rows=1, cols=len(content.columns))
            for i, col in enumerate(content.columns):
                table.cell(0, i).text = col
            for _, row in content.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
        else:
            doc.add_paragraph(content)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def main():
    st.set_page_config(page_title="Mwanza Irrigation Monthly Report", layout="wide")
    st.title("🌾 Mwanza District Irrigation Monthly Report App")

    st.sidebar.header("Upload Images (Optional)")
    uploaded_images = st.sidebar.file_uploader("Upload report images", accept_multiple_files=True, type=["png", "jpg", "jpeg"])

    officer_name = st.text_input("Officer Responsible", "Enter name...")
    report_month = st.date_input("Reporting Month", datetime.today())

    st.subheader("1. Human, Financial and Motor Vehicle Resources")
    staffing = AgGrid(pd.DataFrame(columns=["Post Description", "Grade", "Established", "Filled", "Remarks"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data
    vehicles = AgGrid(pd.DataFrame(columns=["Vehicle Type", "Reg. No.", "Station", "Remarks"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data
    budget = AgGrid(pd.DataFrame(columns=["Cost Centre", "Approved", "Disbursed", "Expenditure", "Percentage", "Remarks"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data

    st.subheader("2. Physical Implementation Progress")
    development = AgGrid(pd.DataFrame(columns=["ID", "Irrigation Technology", "Number of Sites", "Pumps/Canes", "Target Area (ha)", "Actual Cultivated (ha)", "M", "F", "Total Beneficiaries"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data
    utilisation = AgGrid(pd.DataFrame(columns=["EPA", "Area Developed (Ha)", "Target Area (Ha)", "Actual Utilized Area (Ha)", "Beneficiaries"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data

    st.subheader("3. Collaboration with Other Stakeholders")
    stakeholders = AgGrid(pd.DataFrame(columns=["Stakeholder", "Activity"]), editable=True, height=200, update_mode=GridUpdateMode.VALUE_CHANGED).data

    st.subheader("4. Challenges and Recommendations")
    challenges = st.text_area("Challenges", "List main challenges faced...")
    recommendations = st.text_area("Recommendations", "List recommendations...")

    if st.button("Generate Word Report 📄"):
        dataframes = {
            'staffing': staffing,
            'vehicles': vehicles,
            'budget': budget,
            'development': development,
            'utilisation': utilisation,
            'stakeholders': stakeholders,
            'challenges': challenges,
            'recommendations': recommendations
        }
        report = generate_word_report(dataframes, officer_name, report_month)
        st.download_button("📥 Download Report", data=report, file_name="Mwanza_Monthly_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()